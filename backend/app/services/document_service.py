"""
Document service — upload validation, text extraction, chunking, embedding, and CRUD.

All operations that access user data are scoped to a user_id — this is the
multi-tenancy security boundary (same pattern as chat_service).

Processing flow:
  1. validate_file()          — raise ValueError on bad type/size before touching the DB
  2. save_document()          — persist the raw file with status='processing'
  3. process_document()       — extract → chunk → embed → save chunks, then mark ready/failed
     Called as a FastAPI BackgroundTask so the upload endpoint responds immediately.

Supported formats:
  PDF   — pymupdf (fitz): handles text-heavy, LaTeX-generated, and multi-column PDFs.
  DOCX  — python-docx: paragraphs + table cells.
  TXT / MD — decoded as UTF-8 (latin-1 fallback).
  JSON  — parsed and pretty-printed so it's readable in chunks.
  XLSX  — openpyxl: sheets → rows → cells.
  CSV   — stdlib csv module.
  DOC / XLS — legacy formats; accepted at validation but rejected at extraction with a
              clear message asking the user to convert.
"""

from __future__ import annotations

import csv
import io
import json
import logging
from uuid import UUID

import sqlalchemy as sa
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.document import Document, DocumentChunk
from app.services.embedding_service import get_embedding_provider

logger = logging.getLogger(__name__)

# ── Constants ──────────────────────────────────────────────────────────────────

MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20 MB

ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    # DOCX
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    # DOC (legacy Word) — accepted here; extraction will reject with a clear message
    "application/msword",
    "text/plain",
    "text/markdown",
    "application/json",
    "text/json",
    # XLSX
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    # XLS (legacy Excel) — accepted here; extraction will reject with a clear message
    "application/vnd.ms-excel",
    "text/csv",
    "application/csv",
    # Browsers often send these for .md / .txt / .csv files
    "application/octet-stream",
    "text/x-markdown",
}

ALLOWED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".txt", ".md", ".json", ".xlsx", ".xls", ".csv"
}

# ~500 tokens per chunk with ~100 token overlap (4 chars ≈ 1 token heuristic).
_splitter = RecursiveCharacterTextSplitter(
    chunk_size=2000,
    chunk_overlap=400,
    length_function=len,
)

# Embed at most this many chunks per API call to stay under Gemini batch limits.
_EMBED_BATCH_SIZE = 20

# Warn when extraction yields fewer than this many characters — likely a problem.
_SPARSE_TEXT_THRESHOLD = 100


# ── Validation ─────────────────────────────────────────────────────────────────


def validate_file(filename: str, content_type: str, file_size: int) -> None:
    """
    Raise ValueError with a user-facing message if the file is invalid.

    Checks extension, MIME type, and file size — in that order so the most
    specific error is returned first.
    """
    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Allowed: {allowed}"
        )

    # Strip charset suffix (e.g. "text/plain; charset=utf-8") before comparing.
    ct = content_type.lower().split(";")[0].strip()
    if ct not in ALLOWED_CONTENT_TYPES:
        raise ValueError(
            f"Unsupported content type '{content_type}'. "
            f"Upload PDF, DOCX, DOC, TXT, MD, JSON, XLSX, XLS, or CSV files."
        )

    if file_size > MAX_FILE_SIZE_BYTES:
        mb = file_size / (1024 * 1024)
        raise ValueError(f"File too large ({mb:.1f} MB). Maximum allowed size is 20 MB.")


# ── Text extraction ────────────────────────────────────────────────────────────


def extract_text(file_data: bytes, filename: str, content_type: str) -> str:
    """
    Extract plain text from an uploaded file.

    Dispatches primarily by file extension (most reliable), with content_type
    as confirmation. Raises ValueError with a descriptive message on failure.
    """
    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""

    if ext == ".pdf":
        return _extract_pdf(file_data)
    if ext == ".docx":
        return _extract_docx(file_data)
    if ext == ".doc":
        raise ValueError(
            "Legacy .doc format is not supported. "
            "Please re-save the file as .docx (Word 2007+ format) and re-upload."
        )
    if ext in (".txt", ".md"):
        return _extract_plain_text(file_data)
    if ext == ".json":
        return _extract_json(file_data)
    if ext == ".xlsx":
        return _extract_xlsx(file_data)
    if ext == ".xls":
        raise ValueError(
            "Legacy .xls format is not supported. "
            "Please convert the file to .xlsx (Excel 2007+ format) and re-upload."
        )
    if ext == ".csv":
        return _extract_csv(file_data)

    # Fallback for files whose extension slipped past validate_file (shouldn't happen).
    raise ValueError(
        f"Cannot extract text from '{ext}' files. "
        f"Supported: pdf, docx, txt, md, json, xlsx, csv."
    )


def _extract_pdf(file_data: bytes) -> str:
    """Extract text from PDF using pymupdf (fitz). Handles LaTeX and multi-column PDFs."""
    try:
        import fitz  # pymupdf

        doc = fitz.open(stream=file_data, filetype="pdf")
        pages = [page.get_text() for page in doc]
        doc.close()
        text = "\n".join(p for p in pages if p.strip()).strip()
    except Exception as exc:
        raise ValueError(f"PDF extraction failed: {exc}") from exc

    if len(text) < _SPARSE_TEXT_THRESHOLD:
        logger.warning(
            "PDF extraction yielded very little text (%d chars) — "
            "the file may be image-only or contain no selectable text.",
            len(text),
        )
    if not text:
        raise ValueError(
            "PDF contains no extractable text. "
            "It may be image-only or scanned without OCR."
        )
    return text


def _extract_docx(file_data: bytes) -> str:
    """Extract text from DOCX including table cells."""
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_data))
        parts: list[str] = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts).strip()
    except Exception as exc:
        raise ValueError(f"DOCX extraction failed: {exc}") from exc

    if len(text) < _SPARSE_TEXT_THRESHOLD:
        logger.warning("DOCX extraction yielded very little text (%d chars).", len(text))
    if not text:
        raise ValueError("DOCX contains no extractable text.")
    return text


def _extract_plain_text(file_data: bytes) -> str:
    """Decode TXT / MD as UTF-8 with latin-1 fallback for Windows-authored files."""
    try:
        text = file_data.decode("utf-8", errors="replace")
    except Exception:
        text = file_data.decode("latin-1", errors="replace")
    text = text.strip()
    if len(text) < _SPARSE_TEXT_THRESHOLD:
        logger.warning("Plain text extraction yielded very little text (%d chars).", len(text))
    return text


def _extract_json(file_data: bytes) -> str:
    """Parse JSON and re-format with indentation so it reads as structured text."""
    try:
        data = json.loads(file_data.decode("utf-8", errors="replace"))
        text = json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError as exc:
        raise ValueError(f"JSON parsing failed: {exc}") from exc
    if len(text) < _SPARSE_TEXT_THRESHOLD:
        logger.warning("JSON extraction yielded very little text (%d chars).", len(text))
    return text


def _extract_xlsx(file_data: bytes) -> str:
    """Convert an XLSX workbook to readable text: one row per line, cells separated by ' | '."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(
            io.BytesIO(file_data), read_only=True, data_only=True
        )
        parts: list[str] = []
        for sheet in wb.worksheets:
            parts.append(f"=== Sheet: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    parts.append(row_text)
        wb.close()
        text = "\n".join(parts).strip()
    except Exception as exc:
        raise ValueError(f"Excel extraction failed: {exc}") from exc

    if len(text) < _SPARSE_TEXT_THRESHOLD:
        logger.warning("XLSX extraction yielded very little text (%d chars).", len(text))
    if not text:
        raise ValueError("Spreadsheet contains no extractable data.")
    return text


def _extract_csv(file_data: bytes) -> str:
    """Convert CSV rows to readable text: one row per line, cells separated by ' | '."""
    try:
        text_data = file_data.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text_data))
        rows = [
            " | ".join(cell for cell in row)
            for row in reader
            if any(cell.strip() for cell in row)
        ]
        text = "\n".join(rows).strip()
    except Exception as exc:
        raise ValueError(f"CSV extraction failed: {exc}") from exc

    if len(text) < _SPARSE_TEXT_THRESHOLD:
        logger.warning("CSV extraction yielded very little text (%d chars).", len(text))
    if not text:
        raise ValueError("CSV contains no extractable data.")
    return text


# ── Chunking ───────────────────────────────────────────────────────────────────


def chunk_text(text: str) -> list[str]:
    """Split extracted text into overlapping chunks for embedding."""
    return _splitter.split_text(text)


# ── DB operations ──────────────────────────────────────────────────────────────


async def save_document(
    db: AsyncSession,
    user_id: UUID,
    filename: str,
    content_type: str,
    file_data: bytes,
) -> Document:
    """
    Persist the raw document file with status='processing'.

    Chunks are not stored here — they are added by process_document() once
    extraction and embedding have succeeded.
    """
    doc = Document(
        user_id=user_id,
        filename=filename,
        content_type=content_type,
        file_data=file_data,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return doc


async def process_document(
    document_id: UUID,
    user_id: UUID,
    file_data: bytes,
    filename: str,
    content_type: str,
) -> None:
    """
    Background processing step: extract → chunk → embed → store → mark ready.

    Creates its own DB session because this runs after the HTTP response has
    been sent and the request-scoped session is already closed.
    On any failure the document is marked 'failed' with the error message.
    """
    from app.db.session import AsyncSessionLocal  # local import avoids circular dep at module load

    async with AsyncSessionLocal() as db:
        try:
            text = extract_text(file_data, filename, content_type)
            chunks = chunk_text(text)
            if not chunks:
                raise ValueError("No text chunks could be produced from this document.")

            await _save_chunks_with_embeddings(db, document_id, user_id, chunks)
            await _mark_ready(db, document_id)
            logger.info("Document %s processed: %d chunks", document_id, len(chunks))

        except Exception as exc:
            logger.error("Document %s processing failed: %s", document_id, exc)
            try:
                async with AsyncSessionLocal() as err_db:
                    await _mark_failed(err_db, document_id, str(exc))
            except Exception:
                logger.exception("Could not write failure status for document %s", document_id)


async def _save_chunks_with_embeddings(
    db: AsyncSession,
    document_id: UUID,
    user_id: UUID,
    chunks: list[str],
) -> None:
    """
    Embed all chunks in batches and insert them via the ORM.

    Using ORM objects (DocumentChunk) instead of raw SQL avoids the asyncpg
    named-vs-positional parameter mismatch that caused crashes with sa.text().
    pgvector handles the list → vector cast automatically via its SQLAlchemy type.
    """
    embedding_provider = get_embedding_provider()

    for batch_start in range(0, len(chunks), _EMBED_BATCH_SIZE):
        batch = chunks[batch_start : batch_start + _EMBED_BATCH_SIZE]
        vectors = await embedding_provider.embed_texts(batch)

        for offset, (chunk_text_str, vector) in enumerate(zip(batch, vectors)):
            chunk = DocumentChunk(
                document_id=document_id,
                user_id=user_id,
                chunk_text=chunk_text_str,
                embedding=vector,  # plain Python list; pgvector casts to vector(768)
                chunk_index=batch_start + offset,
            )
            db.add(chunk)

    await db.commit()


async def _mark_ready(db: AsyncSession, document_id: UUID) -> None:
    """Set document.status = 'ready' after successful processing."""
    await db.execute(
        sa.update(Document).where(Document.id == document_id).values(status="ready")
    )
    await db.commit()


async def _mark_failed(db: AsyncSession, document_id: UUID, error: str) -> None:
    """Set document.status = 'failed' and store the error message (truncated to 1 KB)."""
    await db.execute(
        sa.update(Document)
        .where(Document.id == document_id)
        .values(status="failed", processing_error=error[:1000])
    )
    await db.commit()


# ── Query operations ───────────────────────────────────────────────────────────


async def list_documents(db: AsyncSession, user_id: UUID) -> list[dict]:
    """
    Return all documents for a user with their chunk counts, newest first.

    Returns dicts with: id, filename, content_type, status, processing_error,
    chunk_count, uploaded_at — never includes file_data (raw bytes).
    """
    result = await db.execute(
        sa.text(
            """
            SELECT
                d.id,
                d.filename,
                d.content_type,
                d.status,
                d.processing_error,
                d.uploaded_at,
                COUNT(c.id)::int AS chunk_count
            FROM documents d
            LEFT JOIN document_chunks c
                ON c.document_id = d.id AND c.user_id = d.user_id
            WHERE d.user_id = :user_id
            GROUP BY d.id
            ORDER BY d.uploaded_at DESC
            """
        ),
        {"user_id": str(user_id)},
    )
    return [dict(row._mapping) for row in result.all()]


async def get_document(
    db: AsyncSession, document_id: UUID, user_id: UUID
) -> Document | None:
    """Return a document only if it exists and belongs to user_id."""
    result = await db.execute(
        sa.select(Document).where(
            Document.id == document_id,
            Document.user_id == user_id,
        )
    )
    return result.scalar_one_or_none()


async def get_document_status(
    db: AsyncSession, document_id: UUID, user_id: UUID
) -> dict | None:
    """
    Return status info for a single document.

    Returns None if the document doesn't exist or isn't owned by user_id.
    """
    result = await db.execute(
        sa.text(
            """
            SELECT
                d.id,
                d.status,
                d.processing_error,
                COUNT(c.id)::int AS chunk_count
            FROM documents d
            LEFT JOIN document_chunks c
                ON c.document_id = d.id AND c.user_id = d.user_id
            WHERE d.id = :document_id AND d.user_id = :user_id
            GROUP BY d.id
            """
        ),
        {"document_id": str(document_id), "user_id": str(user_id)},
    )
    row = result.one_or_none()
    return dict(row._mapping) if row is not None else None


async def delete_document(
    db: AsyncSession, document_id: UUID, user_id: UUID
) -> bool:
    """
    Delete a document and cascade-delete its chunks (via DB ON DELETE CASCADE).

    Returns False if the document doesn't exist or isn't owned by user_id.
    """
    doc = await get_document(db, document_id, user_id)
    if doc is None:
        return False
    await db.delete(doc)
    await db.commit()
    return True
